"""Heading-based document chunking for docx files."""

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


@dataclass
class Section:
    """A document section defined by a heading."""

    heading: str
    level: int
    paragraphs: List[str] = field(default_factory=list)
    tables: List[str] = field(default_factory=list)
    image_refs: List[Dict[str, Any]] = field(default_factory=list)
    page_start: int = 1
    page_end: int = 1
    para_start: int = 0  # paragraph index in document
    para_end: int = 0


class DocxChunker:
    """Splits docx documents into chunks based on heading structure."""

    def __init__(
        self,
        max_tokens: int = 4000,
        overlap_tokens: int = 200,
        strategy: str = "heading",
        image_mode: str = "metadata_only",
    ):
        self.max_tokens = max_tokens
        self.overlap_tokens = overlap_tokens
        self.strategy = strategy
        self.image_mode = image_mode

    def chunk(self, doc_path: str) -> List[Dict[str, Any]]:
        """Parse docx and return list of chunk dictionaries."""
        doc = Document(doc_path)
        filename = Path(doc_path).name
        file_hash = hashlib.md5(filename.encode()).hexdigest()[:8]

        sections = self._build_sections(doc)
        logger.info("Parsed %d sections from %s", len(sections), filename)

        chunks = []
        chunk_index = 0
        for section in sections:
            section_chunks = self._section_to_chunks(section, filename, file_hash, chunk_index)
            chunks.extend(section_chunks)
            chunk_index += len(section_chunks)

        logger.info("Generated %d chunks from %s", len(chunks), filename)
        return chunks

    def _build_sections(self, doc: Document) -> List[Section]:
        """Walk paragraphs and group by heading levels."""
        sections: List[Section] = []
        current_section = Section(heading="(document start)", level=0, para_start=0)
        page_estimate = 1
        para_index = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = Paragraph(element, doc)
                text = para.text.strip()

                # Estimate page breaks
                if self._has_page_break(element):
                    page_estimate += 1

                # Check for images
                images = self._extract_image_refs(element, page_estimate)
                if images:
                    current_section.image_refs.extend(images)

                # Check if this is a heading
                style_name = para.style.name if para.style else ""
                if style_name in HEADING_STYLES and text:
                    # Save current section and start a new one
                    if current_section.paragraphs or current_section.tables:
                        current_section.page_end = page_estimate
                        current_section.para_end = para_index - 1
                        sections.append(current_section)
                    level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                    current_section = Section(
                        heading=text,
                        level=level,
                        page_start=page_estimate,
                        page_end=page_estimate,
                        para_start=para_index,
                    )
                elif text:
                    current_section.paragraphs.append(text)
                    current_section.page_end = page_estimate

                para_index += 1

            elif tag == "tbl":
                table_text = self._table_to_markdown(element, doc)
                if table_text:
                    current_section.tables.append(table_text)

        # Don't forget the last section
        if current_section.paragraphs or current_section.tables:
            current_section.page_end = page_estimate
            current_section.para_end = para_index - 1
            sections.append(current_section)

        # Post-process: infer better heading for "(document start)"
        if sections and sections[0].heading == "(document start)":
            first = sections[0]
            # Use first short paragraph as title-like heading
            for p in first.paragraphs:
                if 2 <= len(p) <= 80:
                    first.heading = p
                    break

        return sections

    def _format_page_range(self, section: Section) -> str:
        """Format page range string, e.g. 'p.1' or 'p.3-7'."""
        if section.page_start == section.page_end:
            return f"p.{section.page_start}"
        return f"p.{section.page_start}-{section.page_end}"

    def _format_position(self, section: Section, part_num: Optional[int] = None) -> str:
        """Format human-readable position string."""
        parts = [f"heading: {section.heading}"]
        parts.append(self._format_page_range(section))
        parts.append(f"para {section.para_start}-{section.para_end}")
        if part_num is not None:
            parts.append(f"part {part_num}")
        return ", ".join(parts)

    def _section_to_chunks(
        self,
        section: Section,
        filename: str,
        file_hash: str,
        start_index: int,
    ) -> List[Dict[str, Any]]:
        """Convert a section into one or more chunks, respecting token limits."""
        # Combine all content
        content_parts = list(section.paragraphs)
        for table_text in section.tables:
            content_parts.append(f"\n[表格]\n{table_text}")

        full_content = "\n".join(content_parts)
        full_tokens = self._estimate_tokens(full_content)

        page_range = self._format_page_range(section)

        # If it fits in one chunk, return as-is
        if full_tokens <= self.max_tokens:
            return [
                self._make_chunk(
                    chunk_id=f"{file_hash}_{start_index}",
                    source_file=filename,
                    source_page=section.page_start,
                    source_page_end=section.page_end,
                    source_section=section.heading,
                    source_position=self._format_position(section),
                    chunk_index=start_index,
                    content=full_content,
                    content_type="text",
                    image_refs=section.image_refs,
                    token_estimate=full_tokens,
                )
            ]

        # Split by paragraphs
        chunks = []
        current_parts: List[str] = []
        current_tokens = 0
        idx = start_index

        for part in content_parts:
            part_tokens = self._estimate_tokens(part)

            if current_tokens + part_tokens > self.max_tokens and current_parts:
                # Flush current chunk
                content = "\n".join(current_parts)
                part_num = idx - start_index + 1
                chunks.append(
                    self._make_chunk(
                        chunk_id=f"{file_hash}_{idx}",
                        source_file=filename,
                        source_page=section.page_start,
                        source_page_end=section.page_end,
                        source_section=section.heading,
                        source_position=self._format_position(section, part_num),
                        chunk_index=idx,
                        content=content,
                        content_type="text",
                        image_refs=section.image_refs if idx == start_index else [],
                        token_estimate=self._estimate_tokens(content),
                    )
                )
                idx += 1

                # Overlap: carry last paragraph
                if self.overlap_tokens > 0 and current_parts:
                    last = current_parts[-1]
                    last_tokens = self._estimate_tokens(last)
                    if last_tokens <= self.overlap_tokens:
                        current_parts = [last]
                        current_tokens = last_tokens
                    else:
                        current_parts = []
                        current_tokens = 0
                else:
                    current_parts = []
                    current_tokens = 0

            current_parts.append(part)
            current_tokens += part_tokens

        # Flush remaining
        if current_parts:
            content = "\n".join(current_parts)
            part_num = idx - start_index + 1
            chunks.append(
                self._make_chunk(
                    chunk_id=f"{file_hash}_{idx}",
                    source_file=filename,
                    source_page=section.page_start,
                    source_page_end=section.page_end,
                    source_section=section.heading,
                    source_position=self._format_position(section, part_num),
                    chunk_index=idx,
                    content=content,
                    content_type="text",
                    image_refs=[],
                    token_estimate=self._estimate_tokens(content),
                )
            )

        return chunks

    def _make_chunk(self, **kwargs) -> Dict[str, Any]:
        """Create a chunk dict with image_refs serialized as JSON string."""
        image_refs = kwargs.get("image_refs", [])
        kwargs["image_refs"] = json.dumps(image_refs, ensure_ascii=False)
        return kwargs

    def _estimate_tokens(self, text: str) -> int:
        """Estimate token count. CJK chars ~0.7 tokens, ASCII words ~1 token."""
        if not text:
            return 0
        cjk_count = sum(1 for c in text if "\u4e00" <= c <= "\u9fff")
        ascii_words = len(re.findall(r"[a-zA-Z0-9]+", text))
        return int(cjk_count * 0.7 + ascii_words + (len(text) - cjk_count) * 0.3)

    def _has_page_break(self, element) -> bool:
        """Check if a paragraph element contains a page break."""
        for br in element.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return False

    def _extract_image_refs(
        self, element, page_estimate: int
    ) -> List[Dict[str, Any]]:
        """Extract image references from a paragraph element."""
        images = []
        for drawing in element.iter(qn("w:drawing")):
            blip = drawing.find(f".//{qn('a:blip')}")
            if blip is not None:
                r_embed = blip.get(qn("r:embed"))
                images.append(
                    {
                        "image_ref": r_embed or "unknown",
                        "page_estimate": page_estimate,
                        "mode": self.image_mode,
                    }
                )
        return images

    def _table_to_markdown(self, tbl_element, doc) -> str:
        """Convert a docx table XML element to markdown-style text."""
        rows = []
        for tr in tbl_element.iter(qn("w:tr")):
            cells = []
            for tc in tr.iter(qn("w:tc")):
                cell_text = ""
                for p in tc.iter(qn("w:p")):
                    para = Paragraph(p, doc)
                    if para.text.strip():
                        cell_text += para.text.strip() + " "
                cells.append(cell_text.strip())
            rows.append(" | ".join(cells))

        if not rows:
            return ""

        # Add header separator after first row
        result = [rows[0]]
        if len(rows) > 1:
            separator = " | ".join(["---"] * len(rows[0].split(" | ")))
            result.append(separator)
            result.extend(rows[1:])

        return "\n".join(result)
