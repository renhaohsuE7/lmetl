"""Shared test fixtures."""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def sample_docx(tmp_path):
    """Create a small docx file with headings, paragraphs, and a table."""
    doc = Document()
    doc.add_heading("第一章 前言", level=1)
    doc.add_paragraph("本報告針對大屯火山群地區進行地熱探勘資料彙整與分析。")
    doc.add_paragraph("研究區域位於臺北市北投區，面積約 50 平方公里。")

    doc.add_heading("1.1 研究目的", level=2)
    doc.add_paragraph("本計畫之主要目的為評估大屯火山群地區之地熱資源潛能。")

    doc.add_heading("第二章 地質概況", level=1)
    doc.add_paragraph("大屯火山群由安山岩及其凝灰岩組成，地質年代為更新世。")
    doc.add_paragraph("主要岩石類型包括安山岩、凝灰角礫�ite及火山碎屑岩。")

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    headers = ["井名", "深度 (m)", "溫度 (°C)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "TH-1"
    table.rows[1].cells[1].text = "1200"
    table.rows[1].cells[2].text = "180"
    table.rows[2].cells[0].text = "TH-2"
    table.rows[2].cells[1].text = "800"
    table.rows[2].cells[2].text = "145"

    doc.add_heading("第三章 結論與建議", level=1)
    doc.add_paragraph("研究結果顯示大屯火山群地區具有良好的地熱資源潛能。")

    path = tmp_path / "test_report.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def mock_llm_response():
    """A mock LLM extraction response."""
    return json.dumps(
        {
            "title": "大屯火山群地區地熱探勘資料執行摘要",
            "authors": ["地質調查所"],
            "institution": "經濟部中央地質調查所",
            "date": "民國109年",
            "year": 2020,
            "abstract": "本報告彙整大屯火山群地區地熱探勘資料。",
            "key_findings": ["大屯火山群具有良好地熱資源潛能"],
            "llm_recommendations": ["建議進行深層鑽探以評估實際產能"],
            "llm_commentary": "此報告為系統性地熱資源評估，資料完整。",
            "confidence_score": 0.85,
            "thinking": "根據文中地質描述與鑽井數據判斷。",
            "rock_types": ["安山岩", "凝灰角礫岩"],
            "formations": ["大屯火山群"],
            "geological_age": "更新世",
            "temperature_gradient": "180°C/1200m",
            "drilling_depth": "1200m",
            "well_names": ["TH-1", "TH-2"],
            "geothermal_assessment": "良好的地熱資源潛能",
            "exploration_methods": ["地球物理探勘", "地球化學分析"],
        },
        ensure_ascii=False,
    )
